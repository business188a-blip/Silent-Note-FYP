"""
exporter.py
Exports meeting minutes to PDF, DOCX and JSON.
Supports Urdu text rendering using NotoNastaliqUrdu font.
"""

import json
import os
from datetime import datetime
from pathlib import Path
from typing import Optional

ASSETS_DIR     = Path(__file__).resolve().parent.parent / "assets"
URDU_FONT_PATH = ASSETS_DIR / "NotoNastaliqUrdu-Regular.ttf"


def _register_urdu_font():
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        if URDU_FONT_PATH.exists():
            pdfmetrics.registerFont(TTFont("NotoNastaliqUrdu", str(URDU_FONT_PATH)))
            return True
    except Exception:
        pass
    return False


class Exporter:

    EXPORTS_DIR = Path(__file__).resolve().parent.parent / "data" / "exports"

    def __init__(self):
        self.EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
        self._urdu_ok = _register_urdu_font()

        try:
            from modules.urdu_helper import fix_urdu, is_urdu
            self._fix_urdu = fix_urdu
            self._is_urdu  = is_urdu
        except Exception:
            self._fix_urdu = lambda x: x
            self._is_urdu  = lambda x: False

    def _smart_para(self, text, normal_style, urdu_style):
        """Return correct Paragraph depending on text language."""
        from reportlab.platypus import Paragraph
        if self._is_urdu(text) and self._urdu_ok:
            fixed = self._fix_urdu(text)
            return Paragraph(fixed, urdu_style)
        return Paragraph(text, normal_style)

    @staticmethod
    def _action_text(item: dict) -> str:
        return item.get("text") or item.get("item") or ""

    def export_pdf(self, meeting: dict, out_path=None) -> str:
        """Generate a PDF file for the given meeting dict."""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles   import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units    import cm
        from reportlab.lib          import colors
        from reportlab.platypus     import (
            SimpleDocTemplate, Paragraph, Spacer,
            Table, TableStyle, HRFlowable, PageBreak,
        )

        out_path = out_path or self._default_path(meeting, "pdf")
        doc = SimpleDocTemplate(
            out_path, pagesize=A4,
            leftMargin=2.5*cm, rightMargin=2.5*cm,
            topMargin=2.5*cm,  bottomMargin=2.5*cm,
        )

        styles    = getSampleStyleSheet()
        urdu_font = "NotoNastaliqUrdu" if self._urdu_ok else "Helvetica"
        story     = []

        h1 = ParagraphStyle("H1SN", parent=styles["Heading1"],
                            fontSize=16, textColor=colors.HexColor("#1a1a2e"),
                            spaceAfter=6)
        h2 = ParagraphStyle("H2SN", parent=styles["Heading2"],
                            fontSize=12, textColor=colors.HexColor("#16213e"),
                            spaceBefore=12, spaceAfter=4)
        body = ParagraphStyle("BodySN", parent=styles["Normal"],
                              fontSize=10, leading=16, spaceAfter=4)

        urdu_style = ParagraphStyle(
            "UrduSN", parent=styles["Normal"],
            fontName  = urdu_font,
            fontSize  = 14,
            leading   = 28,
            spaceAfter= 8,
            alignment = 2,      # right align
        )

        # cover
        story.append(Paragraph(meeting.get("title", "Meeting Minutes"), h1))
        story.append(HRFlowable(width="100%", thickness=1,
                                color=colors.HexColor("#e94560")))
        story.append(Spacer(1, 0.3*cm))

        for label, key in [
            ("Date",         lambda m: self._fmt_dt(m.get("started_at",""))),
            ("Duration",     lambda m: self._fmt_dur(m.get("duration_sec",0))),
            ("Speakers",     lambda m: str(m.get("speaker_count","—"))),
            ("Language",     lambda m: (m.get("language") or "").upper()),
            ("Meeting Mood", lambda m: f"{(m.get('emotion_label') or '—').capitalize()} "
                                       f"({round(float(m.get('emotion_score',0) or 0)*100)}% confidence)"),
        ]:
            val = key(meeting)
            if val and val not in ("—", "0% confidence"):
                story.append(Paragraph(f"<b>{label}:</b> {val}", body))
        story.append(Spacer(1, 0.4*cm))

        # summary
        summary = (meeting.get("summary") or "").strip()
        if summary:
            story.append(Paragraph("Summary", h2))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
            story.append(Spacer(1, 0.2*cm))
            for line in summary.split("\n"):
                if line.strip():
                    story.append(self._smart_para(line.strip(), body, urdu_style))

        # decisions
        decisions = meeting.get("decisions") or []
        if decisions:
            story.append(Paragraph("Key Decisions", h2))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
            story.append(Spacer(1, 0.2*cm))
            for dec in decisions:
                story.append(self._smart_para(f"• {dec}", body, urdu_style))

        # action items
        # FIX: colWidths total = 1 + 9.5 + 3 + 2.5 = 16cm
        # A4 usable width = 21cm - 2.5cm - 2.5cm = 16cm (overflow was causing text bleed)
        action_items = meeting.get("action_items") or []
        if action_items:
            story.append(Paragraph("Action Items", h2))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
            story.append(Spacer(1, 0.2*cm))

            rows = [["#", "Action Item", "Assignee", "Status"]]
            for i, item in enumerate(action_items, 1):
                txt = self._action_text(item)
                if self._is_urdu(txt):
                    txt = self._fix_urdu(txt)
                rows.append([str(i), txt,
                             item.get("assignee","—") or "—",
                             "Done" if item.get("done") else "Pending"])

            # FIX: was [1*cm, 10*cm, 3*cm, 2.5*cm] = 16.5cm — 0.5cm overflow
            tbl = Table(rows, colWidths=[1*cm, 9.5*cm, 3*cm, 2.5*cm])
            tbl.setStyle(TableStyle([
                ("BACKGROUND",    (0,0),(-1,0), colors.HexColor("#16213e")),
                ("TEXTCOLOR",     (0,0),(-1,0), colors.white),
                ("FONTNAME",      (0,0),(-1,0), "Helvetica-Bold"),
                ("FONTSIZE",      (0,0),(-1,0), 9),
                ("ROWBACKGROUNDS",(0,1),(-1,-1),
                 [colors.HexColor("#f5f5f5"), colors.white]),
                ("FONTSIZE",      (0,1),(-1,-1), 9),
                ("GRID",          (0,0),(-1,-1), 0.3, colors.lightgrey),
                ("VALIGN",        (0,0),(-1,-1), "MIDDLE"),
                # FIX: "PADDING" shorthand doesn't exist in ReportLab — use explicit commands
                ("TOPPADDING",    (0,0),(-1,-1), 5),
                ("BOTTOMPADDING", (0,0),(-1,-1), 5),
                ("LEFTPADDING",   (0,0),(-1,-1), 5),
                ("RIGHTPADDING",  (0,0),(-1,-1), 5),
                # FIX: split "FONT" shorthand into separate FONTNAME + FONTSIZE
                ("FONTNAME",      (0,1),(1,-1),  urdu_font),
                ("FONTSIZE",      (0,1),(1,-1),  14),
                ("ALIGN",         (0,1),(1,-1),  "RIGHT"),
            ]))
            story.append(tbl)

        # transcript
        transcript = (meeting.get("transcript") or "").strip()
        if transcript:
            story.append(PageBreak())
            story.append(Paragraph("Full Transcript", h2))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
            story.append(Spacer(1, 0.2*cm))
            for line in transcript.split("\n"):
                if line.strip():
                    story.append(self._smart_para(line.strip(), body, urdu_style))

        doc.build(story)
        return out_path

    def export_docx(self, meeting: dict, out_path=None) -> str:
        """Export as Word document."""
        from docx           import Document
        from docx.shared    import Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        out_path = out_path or self._default_path(meeting, "docx")
        doc      = Document()

        for sec in doc.sections:
            sec.top_margin = sec.bottom_margin = Cm(2.5)
            sec.left_margin = sec.right_margin = Cm(2.5)

        t = doc.add_heading(meeting.get("title","Meeting Minutes"), 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        for key, val in {
            "Date":         self._fmt_dt(meeting.get("started_at","")),
            "Duration":     self._fmt_dur(meeting.get("duration_sec",0)),
            "Speakers":     str(meeting.get("speaker_count","—")),
            "Language":     (meeting.get("language") or "").upper(),
            "Meeting Mood": (meeting.get("emotion_label") or "—").capitalize(),
        }.items():
            p = doc.add_paragraph()
            p.add_run(f"{key}: ").bold = True
            p.add_run(val)

        summary = (meeting.get("summary") or "").strip()
        if summary:
            doc.add_heading("Summary", level=1)
            p = doc.add_paragraph(summary)
            if self._is_urdu(summary):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        for dec in (meeting.get("decisions") or []):
            p = doc.add_paragraph(dec, style="List Bullet")
            if self._is_urdu(dec):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        action_items = meeting.get("action_items") or []
        if action_items:
            doc.add_heading("Action Items", level=1)
            tbl = doc.add_table(rows=1, cols=4)
            tbl.style = "Table Grid"
            for i, h in enumerate(["#","Action Item","Assignee","Status"]):
                tbl.rows[0].cells[i].text = h
                tbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            for idx, item in enumerate(action_items, 1):
                row = tbl.add_row().cells
                row[0].text = str(idx)
                row[1].text = self._action_text(item)
                row[2].text = item.get("assignee","—") or "—"
                row[3].text = "Done" if item.get("done") else "Pending"
                if self._is_urdu(row[1].text):
                    row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        transcript = (meeting.get("transcript") or "").strip()
        if transcript:
            doc.add_page_break()
            doc.add_heading("Full Transcript", level=1)
            for line in transcript.split("\n"):
                if line.strip():
                    p = doc.add_paragraph(line.strip())
                    if self._is_urdu(line):
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.save(out_path)
        return out_path

    def export_json(self, meeting: dict, out_path=None) -> str:
        out_path = out_path or self._default_path(meeting, "json")
        payload  = {
            "exported_at":   datetime.now().isoformat(),
            "title":         meeting.get("title",""),
            "started_at":    meeting.get("started_at",""),
            "ended_at":      meeting.get("ended_at",""),
            "duration_sec":  meeting.get("duration_sec",0),
            "language":      meeting.get("language",""),
            "speaker_count": meeting.get("speaker_count",1),
            "emotion": {
                "label":      meeting.get("emotion_label",""),
                "confidence": meeting.get("emotion_score",0.0),
            },
            "summary":      meeting.get("summary",""),
            "decisions":    meeting.get("decisions",[]),
            "action_items": [
                {"text": self._action_text(i), "assignee": i.get("assignee",""),
                 "due_date": i.get("due_date", ""),
                 "done": bool(i.get("done",False))}
                for i in (meeting.get("action_items") or [])
            ],
            "transcript": meeting.get("transcript",""),
            "snapshots":  [
                {"image_path": s.get("image_path",""),
                 "captured_at": s.get("captured_at",""),
                 "note": s.get("note","")}
                for s in (meeting.get("snapshots") or [])
            ],
        }
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
        return out_path

    def _default_path(self, meeting, ext):
        title     = (meeting.get("title") or "meeting").replace(" ","_")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return str(self.EXPORTS_DIR / f"{title}_{timestamp}.{ext}")

    @staticmethod
    def _fmt_dt(iso):
        if not iso:
            return "—"
        try:
            return datetime.fromisoformat(iso).strftime("%d %B %Y, %I:%M %p")
        except ValueError:
            return iso

    @staticmethod
    def _fmt_dur(seconds):
        try:
            s = int(seconds or 0)
            h = s // 3600
            m = (s % 3600) // 60
            sec = s % 60
            if h:  return f"{h}h {m}m"
            if m:  return f"{m}m {sec}s"
            return f"{sec}s"
        except Exception:
            return "0s"